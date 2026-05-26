#!/usr/bin/env python3
"""
PaperForge Word Guard
Checks a paper.docx file for common structural and content issues.
Usage:
    python word_guard.py paper_forge_output/final_paper/paper.docx --markdown
"""

from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass
from pathlib import Path


@dataclass
class WordCheck:
    name: str
    passed: bool
    detail: str = ""


def check_word(docx_path: Path) -> list[WordCheck]:
    checks = []

    def add(name: str, condition: bool, detail: str = ""):
        checks.append(WordCheck(name=name, passed=condition, detail=detail))

    try:
        from docx import Document
        doc = Document(str(docx_path))
    except ImportError:
        add("python-docx installed", False, "pip install python-docx")
        return checks
    except Exception as e:
        add("File readable", False, str(e))
        return checks

    add("File readable", True)

    # Paragraph count
    paras = [p for p in doc.paragraphs if p.text.strip()]
    add("Has paragraphs", len(paras) > 0, f"{len(paras)} non-empty paragraphs")
    add("Minimum content", len(paras) >= 5, f"Found {len(paras)}, expected ≥ 5")

    # Heading check
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    add("Has headings", len(headings) > 0, f"{len(headings)} heading(s) found")

    # Word count estimate
    all_text = " ".join(p.text for p in doc.paragraphs)
    word_count = len(all_text.split())
    add("Word count ≥ 500", word_count >= 500, f"Estimated {word_count} words")

    # Check for placeholder text
    placeholder_markers = ["[PLACEHOLDER]", "TODO", "FIXME", "INSERT HERE", "TBD"]
    found_placeholders = [m for m in placeholder_markers if m.lower() in all_text.lower()]
    add(
        "No placeholder text",
        len(found_placeholders) == 0,
        f"Found: {found_placeholders}" if found_placeholders else "",
    )

    # Tables
    table_count = len(doc.tables)
    add("Tables check", True, f"{table_count} table(s) present")

    # Sections / styles used
    styles_used = {p.style.name for p in doc.paragraphs if p.text.strip()}
    add("Multiple styles", len(styles_used) > 1, f"Styles used: {', '.join(sorted(styles_used)[:5])}")

    return checks


def format_markdown(checks: list[WordCheck], docx_path: Path) -> str:
    passed = sum(1 for c in checks if c.passed)
    total = len(checks)
    lines = [
        "# PaperForge Word Guard\n",
        f"**File:** `{docx_path}`\n",
        f"**Checks passed:** {passed}/{total}\n",
        "",
        "| Status | Check | Detail |",
        "|--------|-------|--------|",
    ]
    for c in checks:
        icon = "✅" if c.passed else "❌"
        lines.append(f"| {icon} | {c.name} | {c.detail} |")

    failed = [c for c in checks if not c.passed]
    if failed:
        lines += ["", "## Issues to Fix\n"]
        for c in failed:
            detail = f": {c.detail}" if c.detail else ""
            lines.append(f"- **{c.name}**{detail}")

    overall = "**PASS**" if not failed else "**FAIL**"
    lines += ["", f"## Overall: {overall}\n"]
    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(description="PaperForge Word Guard")
    parser.add_argument("docx_file", help="Path to paper.docx")
    parser.add_argument("--markdown", action="store_true")
    parser.add_argument("--write", action="store_true")
    args = parser.parse_args()

    docx_path = Path(args.docx_file)
    if not docx_path.exists():
        print(f"Error: file not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    checks = check_word(docx_path)
    failed = [c for c in checks if not c.passed]

    if args.markdown:
        report = format_markdown(checks, docx_path)
        print(report)
        if args.write:
            out = docx_path.parent / "word_guard.md"
            out.write_text(report, encoding="utf-8")
            print(f"\nReport written to: {out}")
    else:
        for c in checks:
            icon = "OK  " if c.passed else "FAIL"
            print(f"  [{icon}] {c.name}" + (f" — {c.detail}" if c.detail else ""))

    sys.exit(1 if failed else 0)


if __name__ == "__main__":
    main()
